#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
file_parser.py - Parse various file formats (docx, pdf, txt, md) to extract text content

Supports:
- .docx files (using python-docx)
- .pdf files (using PyPDF2 or pdfplumber)
- .txt files (plain text)
- .md files (markdown)
"""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from typing import Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: Path) -> Optional[str]:
    """
    Extract text content from various file formats
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content, or None if extraction fails
        
    Supported formats:
        - .docx: Microsoft Word documents
        - .pdf: PDF documents
        - .txt: Plain text files
        - .md: Markdown files
    """
    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        return None
    
    file_ext = file_path.suffix.lower()
    
    try:
        if file_ext == '.docx':
            return _extract_text_from_docx(file_path)
        elif file_ext == '.pdf':
            return _extract_text_from_pdf(file_path)
        elif file_ext in ['.txt', '.md']:
            return _extract_text_from_text(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return None
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}", exc_info=True)
        return None


def _extract_text_from_docx(file_path: Path) -> Optional[str]:
    """Extract text from .docx file using python-docx"""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Install with: pip install python-docx")
        return None
    
    try:
        doc = Document(file_path)
        paragraphs = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}", exc_info=True)
        return None


def _extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """Extract text from .pdf file using PyPDF2 or pdfplumber"""
    # Try pdfplumber first (better for complex PDFs)
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
    except ImportError:
        # Fallback to PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                pages = []
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                return "\n\n".join(pages)
        except ImportError:
            logger.error(
                "Neither pdfplumber nor PyPDF2 is installed. "
                "Install with: pip install pdfplumber (recommended) or pip install PyPDF2"
            )
            return None
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}", exc_info=True)
        return None


def _extract_text_from_text(file_path: Path) -> Optional[str]:
    """Extract text from plain text or markdown file"""
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except UnicodeDecodeError:
        # Fallback to latin-1
        try:
            with open(file_path, 'r', encoding='latin-1', errors='replace') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}", exc_info=True)
            return None


def get_file_metadata(file_path: Path) -> dict:
    """
    Get metadata for a file
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file metadata (size, type, etc.)
    """
    metadata = {
        'file_name': file_path.name,
        'file_size': file_path.stat().st_size if file_path.exists() else 0,
        'file_type': file_path.suffix.lower(),
        'exists': file_path.exists(),
    }
    
    # Add format-specific metadata
    if file_path.suffix.lower() == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
        except:
            pass
    elif file_path.suffix.lower() == '.pdf':
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                metadata['page_count'] = len(pdf.pages)
        except:
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)
            except:
                pass
    
    return metadata


if __name__ == "__main__":
    # Test file parsing
    import argparse
    
    parser = argparse.ArgumentParser(description="Extract text from files")
    parser.add_argument("file_path", type=Path, help="Path to file")
    parser.add_argument("--metadata", action="store_true", help="Show file metadata")
    
    args = parser.parse_args()
    
    if args.metadata:
        metadata = get_file_metadata(args.file_path)
        print("File Metadata:")
        for key, value in metadata.items():
            print(f"  {key}: {value}")
    else:
        text = extract_text_from_file(args.file_path)
        if text:
            print("Extracted Text:")
            print("=" * 80)
            print(text[:1000])  # Print first 1000 characters
            if len(text) > 1000:
                print(f"\n... (truncated, total length: {len(text)} characters)")
        else:
            print("Failed to extract text from file")
