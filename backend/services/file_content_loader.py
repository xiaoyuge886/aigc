"""
File content loader for loading file contents from docs-management or file system

This service supports:
1. Query-based retrieval: Search file content by keywords or semantic query
2. Meta information usage: Leverage file metadata for intelligent retrieval
3. Fallback strategy: Use keyword-based local search if docs-management is unavailable
"""
from pathlib import Path
from typing import Optional, List, Dict, Any

from loguru import logger

from services.database import DatabaseService
from models.database import UserFileRelationshipDB


class FileContent:
    """File content object"""
    def __init__(
        self,
        doc_id: str,
        file_name: str,
        content: str,
        file_type: str,
        section: Optional[str] = None,
        relevance_score: float = 1.0,
    ):
        self.doc_id = doc_id
        self.file_name = file_name
        self.content = content
        self.file_type = file_type
        self.section = section
        self.relevance_score = relevance_score


class FileContentSnippet:
    """File content snippet from search result"""
    def __init__(
        self,
        doc_id: str,
        file_name: str,
        snippet: str,
        file_type: str,
        relevance_score: float = 0.0,
        line_start: Optional[int] = None,
        line_end: Optional[int] = None,
    ):
        self.doc_id = doc_id
        self.file_name = file_name
        self.snippet = snippet
        self.file_type = file_type
        self.relevance_score = relevance_score
        self.line_start = line_start
        self.line_end = line_end


class FileContentLoader:
    """
    Service for loading file contents with intelligent retrieval
    
    Features:
    - Query-based retrieval: Search content by keywords or semantic query
    - Meta information usage: Leverage file metadata for intelligent retrieval
    - Fallback strategy: Use keyword-based local search if docs-management unavailable
    """
    
    def __init__(self, db_service: DatabaseService):
        self.db_service = db_service
        # Base directory for user uploads
        from core.config import settings
        self.uploads_base_dir = Path(settings.work_dir) / ".claude" / "skills" / "docs-management" / "canonical" / "user-uploads"
        # Try to initialize docs-management integration (if available)
        self._docs_management_available = self._check_docs_management_available()
    
    def _check_docs_management_available(self) -> bool:
        """Check if docs-management skill is available"""
        try:
            from core.config import settings
            docs_management_dir = Path(settings.work_dir) / ".claude" / "skills" / "docs-management"
            # Check if docs-management skill directory exists
            if docs_management_dir.exists():
                # Try to import docs-management API if available
                # For now, we'll use a simple check
                logger.info("docs-management skill directory found, but API integration not yet implemented")
                return False
            return False
        except Exception as e:
            logger.debug(f"docs-management not available: {e}")
            return False
    
    async def load_file_content(
        self,
        doc_id: str,
        section: Optional[str] = None,
        user_id: Optional[int] = None
    ) -> FileContent:
        """
        Load file content
        
        Args:
            doc_id: File doc_id
            section: Optional subsection name to extract
            user_id: User ID for permission verification
        
        Returns:
            FileContent - File content object
        
        Raises:
            PermissionError: If user doesn't have permission
            FileNotFoundError: If file doesn't exist
        """
        # 1. Verify permission
        relationship = await self.db_service.get_file_relationship(doc_id)
        if not relationship:
            raise FileNotFoundError(f"File not found: {doc_id}")
        
        if user_id and relationship.user_id != user_id:
            raise PermissionError(f"Permission denied: User {user_id} cannot access file {doc_id}")
        
        # 2. Get file path from relationship
        # For now, we reconstruct the path from doc_id
        # In the future, we could store the path in the relationship or get it from docs-management
        user_id_from_doc = doc_id.split('-')[2]  # Extract user_id from doc_id
        file_path = self._get_file_path_from_doc_id(doc_id, relationship.file_name)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found at path: {file_path}")
        
        # 3. Load file content
        if section:
            # TODO: Extract subsection (requires integration with docs-management)
            # For now, load full content
            logger.warning(f"Subsection extraction not yet implemented, loading full content")
            content = self._load_full_content(file_path, relationship.file_type)
        else:
            content = self._load_full_content(file_path, relationship.file_type)
        
        return FileContent(
            doc_id=doc_id,
            file_name=relationship.file_name,
            content=content,
            file_type=relationship.file_type,
            section=section
        )
    
    def _get_file_path_from_doc_id(self, doc_id: str, file_name: str) -> Path:
        """Get file path from doc_id"""
        # Extract user_id from doc_id (format: user-upload-{user_id}-{hash})
        parts = doc_id.split('-')
        if len(parts) < 4:
            raise ValueError(f"Invalid doc_id format: {doc_id}")
        
        user_id = parts[2]
        file_extension = Path(file_name).suffix
        saved_file_name = f"{doc_id}{file_extension}"
        
        return self.uploads_base_dir / user_id / saved_file_name
    
    def get_file_path(self, doc_id: str, file_name: str) -> Path:
        """Get file path from doc_id (public method)"""
        return self._get_file_path_from_doc_id(doc_id, file_name)
    
    def _load_full_content(self, file_path: Path, file_type: str) -> str:
        """Load full file content"""
        try:
            # For text files, read directly
            # Also check file extension as fallback (some files may have incorrect MIME type)
            text_extensions = {'.txt', '.md', '.markdown', '.json', '.xml', '.yaml', '.yml', '.csv', '.log'}
            parseable_extensions = {'.docx', '.pdf'}  # Can be parsed to text
            
            is_text_file = (
                file_type.startswith('text/') or 
                file_type in [
                    'application/json',
                    'application/xml',
                    'application/x-yaml',
                    'text/markdown',
                    'text/plain',
                ] or
                file_path.suffix.lower() in text_extensions
            )
            is_parseable = file_path.suffix.lower() in parseable_extensions
            
            if is_text_file:
                return file_path.read_text(encoding='utf-8')
            elif is_parseable:
                # For docx/pdf files, extract text using _extract_text_from_file
                text_content = self._extract_text_from_file(file_path, file_type)
                if text_content:
                    return text_content
                else:
                    return f"[无法提取文本内容: {file_path.name} (类型: {file_type})]"
            elif file_type.startswith('image/'):
                return f"[图片文件: {file_path.name}]"
            else:
                # Try to read as text (might work for some files)
                try:
                    return file_path.read_text(encoding='utf-8')
                except:
                    return f"[无法读取文件内容: {file_path.name} (类型: {file_type})]"
        except Exception as e:
            logger.error(f"Failed to load file content from {file_path}: {e}", exc_info=True)
            raise
    
    async def get_file_summary(
        self,
        doc_id: str,
        max_lines: int = 50,
        user_id: Optional[int] = None
    ) -> str:
        """
        Get file summary (first N lines) instead of full content
        This is used to avoid loading large files into prompt
        
        Args:
            doc_id: File doc_id
            max_lines: Maximum number of lines to include in summary
            user_id: User ID for permission verification
        
        Returns:
            str: File summary (first N lines)
        """
        # 1. Verify permission
        relationship = await self.db_service.get_file_relationship(doc_id)
        if not relationship:
            raise FileNotFoundError(f"File not found: {doc_id}")
        
        if user_id and relationship.user_id != user_id:
            raise PermissionError(f"Permission denied: User {user_id} cannot access file {doc_id}")
        
        # 2. Get file path
        file_path = self._get_file_path_from_doc_id(doc_id, relationship.file_name)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found at path: {file_path}")
        
        # 3. Load summary (first N lines)
        try:
            text_extensions = {'.txt', '.md', '.markdown', '.json', '.xml', '.yaml', '.yml', '.csv', '.log'}
            parseable_extensions = {'.docx', '.pdf'}  # Can be parsed to text
            
            is_text_file = (
                relationship.file_type.startswith('text/') or 
                relationship.file_type in [
                    'application/json',
                    'application/xml',
                    'application/x-yaml',
                    'text/markdown',
                    'text/plain',
                ] or
                file_path.suffix.lower() in text_extensions
            )
            is_parseable = file_path.suffix.lower() in parseable_extensions
            
            if is_text_file:
                lines = file_path.read_text(encoding='utf-8').split('\n')
                total_lines = len(lines)
                summary_lines = lines[:max_lines]
                summary = '\n'.join(summary_lines)
                
                if total_lines > max_lines:
                    summary += f"\n\n... (文件共 {total_lines} 行，此处仅显示前 {max_lines} 行。如需查看完整内容，请使用 Read 工具读取文件)"
                
                return summary
            elif is_parseable:
                # For docx/pdf files, extract text and return summary
                text_content = self._extract_text_from_file(file_path, relationship.file_type)
                if text_content:
                    lines = text_content.split('\n')
                    total_lines = len(lines)
                    summary_lines = lines[:max_lines]
                    summary = '\n'.join(summary_lines)
                    
                    if total_lines > max_lines:
                        summary += f"\n\n... (文件共 {total_lines} 行，此处仅显示前 {max_lines} 行。如需查看完整内容，请使用 Read 工具读取文件)"
                    
                    return summary
                else:
                    return f"[无法提取文本内容。文件类型: {relationship.file_type}, 大小: {relationship.file_size} 字节]"
            else:
                # For non-text and non-parseable files, return file info
                return f"[文件类型: {relationship.file_type}, 大小: {relationship.file_size} 字节。请使用适当的工具读取此文件]"
        except Exception as e:
            logger.error(f"Failed to get file summary from {file_path}: {e}", exc_info=True)
            raise
    
    async def search_file_content(
        self,
        doc_id: str,
        query: str,
        user_id: Optional[int] = None,
        max_snippets: int = 5,
        context_lines: int = 3,
    ) -> List[FileContentSnippet]:
        """
        Search file content by query (keywords or semantic search)
        
        This method uses intelligent retrieval instead of loading the entire file:
        1. First tries docs-management semantic search (if available)
        2. Falls back to keyword-based local search
        3. Returns relevant snippets with relevance scores
        
        Args:
            doc_id: File doc_id
            query: Search query (keywords or natural language)
            user_id: User ID for permission verification
            max_snippets: Maximum number of snippets to return
            context_lines: Number of context lines around each match
        
        Returns:
            List[FileContentSnippet] - List of relevant content snippets
        """
        # 1. Verify permission
        relationship = await self.db_service.get_file_relationship(doc_id)
        if not relationship:
            raise FileNotFoundError(f"File not found: {doc_id}")
        
        if user_id and relationship.user_id != user_id:
            raise PermissionError(f"Permission denied: User {user_id} cannot access file {doc_id}")
        
        # 2. Try docs-management search first (if available)
        if self._docs_management_available:
            try:
                snippets = await self._search_via_docs_management(
                    doc_id=doc_id,
                    query=query,
                    max_snippets=max_snippets,
                    context_lines=context_lines,
                )
                if snippets:
                    logger.info(f"Found {len(snippets)} snippets via docs-management for query: {query}")
                    return snippets
            except Exception as e:
                logger.warning(f"docs-management search failed, falling back to local search: {e}")
        
        # 3. Fallback to keyword-based local search
        return await self._search_by_keywords_local(
            doc_id=doc_id,
            query=query,
            relationship=relationship,
            max_snippets=max_snippets,
            context_lines=context_lines,
        )
    
    async def _search_via_docs_management(
        self,
        doc_id: str,
        query: str,
        max_snippets: int,
        context_lines: int,
    ) -> List[FileContentSnippet]:
        """
        Search via docs-management skill (if available)
        
        TODO: Implement integration with docs-management skill API
        """
        # Placeholder for docs-management integration
        # This would call docs-management's search API
        logger.debug(f"docs-management search not yet implemented for doc_id: {doc_id}, query: {query}")
        return []
    
    async def _search_by_keywords_local(
        self,
        doc_id: str,
        query: str,
        relationship: UserFileRelationshipDB,
        max_snippets: int,
        context_lines: int,
    ) -> List[FileContentSnippet]:
        """
        Local keyword-based search as fallback
        
        This method:
        1. Extracts keywords from query
        2. Searches file content for matching lines
        3. Returns snippets with context around matches
        4. Calculates relevance scores based on keyword matches
        """
        # 1. Get file path
        file_path = self._get_file_path_from_doc_id(doc_id, relationship.file_name)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found at path: {file_path}")
        
        # 2. Check if file is text-based or can be parsed (docx, pdf)
        text_extensions = {'.txt', '.md', '.markdown', '.json', '.xml', '.yaml', '.yml', '.csv', '.log', '.py', '.js', '.ts', '.html', '.css'}
        parseable_extensions = {'.docx', '.pdf'}  # Can be parsed to text
        is_text_file = (
            relationship.file_type.startswith('text/') or 
            relationship.file_type in [
                'application/json',
                'application/xml',
                'application/x-yaml',
                'text/markdown',
                'text/plain',
            ] or
            file_path.suffix.lower() in text_extensions
        )
        is_parseable = file_path.suffix.lower() in parseable_extensions
        
        if not is_text_file and not is_parseable:
            # For non-text and non-parseable files, return file info
            return [FileContentSnippet(
                doc_id=doc_id,
                file_name=relationship.file_name,
                snippet=f"[文件类型: {relationship.file_type}, 大小: {relationship.file_size} 字节。请使用适当的工具读取此文件]",
                file_type=relationship.file_type,
                relevance_score=0.5,
            )]
        
        # 3. Extract keywords from query
        keywords = self._extract_keywords(query)
        if not keywords:
            # If no keywords, return file summary
            summary = await self.get_file_summary(doc_id, max_lines=50)
            return [FileContentSnippet(
                doc_id=doc_id,
                file_name=relationship.file_name,
                snippet=summary,
                file_type=relationship.file_type,
                relevance_score=0.3,
            )]
        
        # 4. Read file and search for keywords
        try:
            # Use file parser for docx/pdf support
            text_content = self._extract_text_from_file(file_path, relationship.file_type)
            if not text_content:
                # Fallback: try reading as text
                try:
                    text_content = file_path.read_text(encoding='utf-8')
                except:
                    logger.warning(f"Could not extract text from {file_path}")
                    return [FileContentSnippet(
                        doc_id=doc_id,
                        file_name=relationship.file_name,
                        snippet=f"[无法提取文本内容。文件类型: {relationship.file_type}, 大小: {relationship.file_size} 字节]",
                        file_type=relationship.file_type,
                        relevance_score=0.1,
                    )]
            
            lines = text_content.split('\n')
            snippets = self._find_relevant_snippets(
                lines=lines,
                keywords=keywords,
                max_snippets=max_snippets,
                context_lines=context_lines,
            )
            
            # 5. Create FileContentSnippet objects
            result = []
            for snippet_data in snippets:
                result.append(FileContentSnippet(
                    doc_id=doc_id,
                    file_name=relationship.file_name,
                    snippet=snippet_data['snippet'],
                    file_type=relationship.file_type,
                    relevance_score=snippet_data['score'],
                    line_start=snippet_data.get('line_start'),
                    line_end=snippet_data.get('line_end'),
                ))
            
            logger.info(f"Found {len(result)} relevant snippets for query '{query}' in file {relationship.file_name}")
            return result
            
        except Exception as e:
            logger.error(f"Failed to search file content: {e}", exc_info=True)
            raise
    
    def _extract_keywords(self, query: str) -> List[str]:
        """Extract keywords from query"""
        # Simple keyword extraction: split by spaces and filter
        # In the future, this could use NLP for better extraction
        keywords = []
        for word in query.split():
            word = word.strip().lower()
            # Filter out common stop words (simple list)
            stop_words = {'的', '是', '在', '有', '和', '与', '或', '但', '如果', 'the', 'a', 'an', 'and', 'or', 'but', 'if', 'is', 'are', 'was', 'were'}
            if len(word) > 1 and word not in stop_words:
                keywords.append(word)
        return keywords
    
    def _find_relevant_snippets(
        self,
        lines: List[str],
        keywords: List[str],
        max_snippets: int,
        context_lines: int,
    ) -> List[Dict[str, Any]]:
        """
        Find relevant snippets in file content
        
        Returns snippets with:
        - Matching lines with context
        - Relevance scores based on keyword matches
        """
        matches = []
        
        # Find lines containing keywords
        for i, line in enumerate(lines):
            line_lower = line.lower()
            # Count keyword matches in this line
            match_count = sum(1 for keyword in keywords if keyword in line_lower)
            if match_count > 0:
                # Calculate relevance score (normalized by number of keywords)
                score = match_count / len(keywords)
                matches.append({
                    'line_index': i,
                    'line': line,
                    'score': score,
                    'match_count': match_count,
                })
        
        # Sort by relevance score (descending)
        matches.sort(key=lambda x: x['score'], reverse=True)
        
        # Take top matches
        top_matches = matches[:max_snippets]
        
        # Build snippets with context
        snippets = []
        used_lines = set()
        
        for match in top_matches:
            line_idx = match['line_index']
            
            # Skip if this line is already included in another snippet
            if line_idx in used_lines:
                continue
            
            # Calculate snippet range
            start_line = max(0, line_idx - context_lines)
            end_line = min(len(lines), line_idx + context_lines + 1)
            
            # Mark lines as used
            for i in range(start_line, end_line):
                used_lines.add(i)
            
            # Build snippet text
            snippet_lines = []
            for i in range(start_line, end_line):
                prefix = ">>> " if i == line_idx else "    "
                snippet_lines.append(f"{prefix}第 {i+1} 行: {lines[i]}")
            
            snippet_text = '\n'.join(snippet_lines)
            
            snippets.append({
                'snippet': snippet_text,
                'score': match['score'],
                'line_start': start_line + 1,  # 1-indexed
                'line_end': end_line,
            })
        
        return snippets
    
    async def get_file_content_by_query(
        self,
        doc_id: str,
        query: str,
        user_id: Optional[int] = None,
        max_length: int = 5000,
    ) -> str:
        """
        Get file content relevant to a query
        
        This is a convenience method that:
        1. Uses search_file_content to find relevant snippets
        2. Combines snippets into a single string
        3. Limits total length to avoid token overflow
        
        Args:
            doc_id: File doc_id
            query: Search query
            user_id: User ID for permission verification
            max_length: Maximum total length of returned content
        
        Returns:
            str: Relevant file content snippets
        """
        snippets = await self.search_file_content(
            doc_id=doc_id,
            query=query,
            user_id=user_id,
            max_snippets=10,  # Get more snippets, then trim
        )
        
        if not snippets:
            # Fallback to summary
            return await self.get_file_summary(doc_id, max_lines=50, user_id=user_id)
        
        # Combine snippets, sorted by relevance
        snippets.sort(key=lambda x: x.relevance_score, reverse=True)
        
        result_parts = []
        total_length = 0
        
        for snippet in snippets:
            snippet_text = f"\n\n--- 相关片段 (相关性: {snippet.relevance_score:.2f}) ---\n{snippet.snippet}"
            snippet_length = len(snippet_text)
            
            if total_length + snippet_length > max_length:
                # Add partial snippet if there's room
                remaining = max_length - total_length
                if remaining > 100:  # Only add if meaningful space remains
                    result_parts.append(snippet_text[:remaining] + "\n... (内容已截断)")
                break
            
            result_parts.append(snippet_text)
            total_length += snippet_length
        
        if not result_parts:
            # Fallback to summary
            return await self.get_file_summary(doc_id, max_lines=50, user_id=user_id)
        
        return "".join(result_parts)
    
    def _extract_text_from_file(self, file_path: Path, file_type: str) -> Optional[str]:
        """
        Extract text from various file formats (docx, pdf, txt, md)
        
        Args:
            file_path: Path to the file
            file_type: MIME type of the file
            
        Returns:
            Extracted text content, or None if extraction fails
        """
        if not file_path.exists():
            return None
        
        file_ext = file_path.suffix.lower()
        
        # Try to use docs-management file parser if available
        try:
            from pathlib import Path as PathLib
            import sys
            docs_management_dir = Path(__file__).parent.parent.parent / ".claude" / "skills" / "docs-management"
            if docs_management_dir.exists():
                sys.path.insert(0, str(docs_management_dir / "scripts"))
                from utils.file_parser import extract_text_from_file
                return extract_text_from_file(file_path)
        except ImportError:
            logger.debug("docs-management file parser not available, using fallback")
        except Exception as e:
            logger.warning(f"Failed to use docs-management file parser: {e}")
        
        # Fallback: handle common file types directly
        try:
            if file_ext == '.docx':
                # Try python-docx
                try:
                    from docx import Document
                    doc = Document(file_path)
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    # Extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                paragraphs.append(" | ".join(row_text))
                    return "\n\n".join(paragraphs)
                except ImportError:
                    logger.warning("python-docx not installed. Install with: pip install python-docx")
                    return None
            
            elif file_ext == '.pdf':
                # Try pdfplumber first (better quality)
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        pages = []
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                pages.append(text)
                        return "\n\n".join(pages)
                except ImportError:
                    # Fallback to PyPDF2
                    try:
                        import PyPDF2
                        with open(file_path, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            pages = []
                            for page_num in range(len(pdf_reader.pages)):
                                page = pdf_reader.pages[page_num]
                                text = page.extract_text()
                                if text:
                                    pages.append(text)
                            return "\n\n".join(pages)
                    except ImportError:
                        logger.warning(
                            "Neither pdfplumber nor PyPDF2 is installed. "
                            "Install with: pip install pdfplumber (recommended) or pip install PyPDF2"
                        )
                        return None
            
            elif file_ext in ['.txt', '.md', '.markdown']:
                # Plain text files
                try:
                    return file_path.read_text(encoding='utf-8', errors='replace')
                except:
                    return file_path.read_text(encoding='latin-1', errors='replace')
            
            else:
                # Unknown file type, try reading as text
                try:
                    return file_path.read_text(encoding='utf-8', errors='replace')
                except:
                    return None
                    
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}", exc_info=True)
            return None
